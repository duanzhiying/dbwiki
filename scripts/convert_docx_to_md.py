#!/usr/bin/env python3
"""
将DOCX文件转换为Markdown格式
"""

import docx
import os
from pathlib import Path


def convert_docx_to_md(docx_path, output_path):
    """将DOCX文件转换为Markdown格式"""
    try:
        # 打开DOCX文件
        doc = docx.Document(docx_path)
        
        # 准备Markdown内容
        md_content = []
        
        # 处理每个段落
        for para in doc.paragraphs:
            if para.text.strip():
                # 检查是否为标题（根据缩进和格式判断）
                if para.style.name.startswith('Heading'):
                    level = int(para.style.name[-1])
                    md_content.append('#' * level + ' ' + para.text)
                else:
                    md_content.append(para.text)
        
        # 处理表格
        for table in doc.tables:
            table_md = []
            # 处理表头
            header_cells = table.rows[0].cells
            table_md.append('| ' + ' | '.join(cell.text for cell in header_cells) + ' |')
            table_md.append('| ' + ' | '.join(['---' for _ in header_cells]) + ' |')
            # 处理数据行
            for row in table.rows[1:]:
                table_md.append('| ' + ' | '.join(cell.text for cell in row.cells) + ' |')
            md_content.extend(table_md)
        
        # 保存为Markdown文件
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n\n'.join(md_content))
        
        print(f"✅ 转换完成：{docx_path} -> {output_path}")
        return True
        
    except Exception as e:
        print(f"❌ 转换失败：{str(e)}")
        return False


def main():
    """主函数"""
    docx_file = Path('/Users/kidzying/Documents/紫金学院/教改项目实施/dbwiki/knowledge_base/raw_docs/mysql  workbench的打开.docx')
    output_file = Path('/Users/kidzying/Documents/紫金学院/教改项目实施/dbwiki/knowledge_base/raw_docs/mysql_workbench_install.md')
    
    if not docx_file.exists():
        print(f"❌ 文件不存在：{docx_file}")
        return 1
    
    success = convert_docx_to_md(docx_file, output_file)
    return 0 if success else 1


if __name__ == "__main__":
    exit(main())
